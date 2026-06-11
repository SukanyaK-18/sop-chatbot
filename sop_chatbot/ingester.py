from __future__ import annotations

import base64
import io
import os

from langchain_text_splitters import RecursiveCharacterTextSplitter

from sop_chatbot.config import ChatbotConfig
from sop_chatbot.index import SOPIndex
from sop_chatbot.models import Chunk, IngestError, IngestResult, DocumentNotFoundError

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx"}

# Directory to store extracted images
IMAGES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "sop_images")


def _save_images_from_docx(file_path: str) -> list[str]:
    """Extract all images from a .docx file, save to disk, return file paths."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    os.makedirs(IMAGES_DIR, exist_ok=True)

    file_base = os.path.splitext(os.path.basename(file_path))[0]
    image_paths: list[str] = []
    idx = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_blob = rel.target_part.blob
            # Determine extension from content type
            content_type = rel.target_part.content_type
            if "png" in content_type:
                ext = ".png"
            elif "jpeg" in content_type or "jpg" in content_type:
                ext = ".jpg"
            elif "gif" in content_type:
                ext = ".gif"
            else:
                ext = ".png"

            img_filename = f"{file_base}_image_{idx}{ext}"
            img_path = os.path.join(IMAGES_DIR, img_filename)
            with open(img_path, "wb") as f:
                f.write(image_blob)
            image_paths.append(img_path)
            idx += 1

    return image_paths


def _extract_images_from_docx(file_path: str) -> list[str]:
    """Extract all images from a .docx file as base64-encoded strings."""
    from docx import Document as DocxDocument
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = DocxDocument(file_path)
    images: list[str] = []

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_blob = rel.target_part.blob
            b64 = base64.b64encode(image_blob).decode("utf-8")
            images.append(b64)

    return images


def _describe_image_with_vision(image_b64: str, image_index: int) -> str:
    """Send a base64 image to Groq's vision model and get a text description."""
    import groq as groq_sdk

    try:
        client = groq_sdk.Groq()
        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            max_tokens=1024,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Describe this image from an SOP document in detail. "
                                "Include all text, labels, steps, arrows, and any information visible. "
                                "Format your response as plain text that captures all the information in the image."
                            ),
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{image_b64}",
                            },
                        },
                    ],
                }
            ],
        )
        description = response.choices[0].message.content
        return f"[Image {image_index + 1} description]: {description}"
    except Exception:
        return f"[Image {image_index + 1}: unable to process]"


def _extract_docx_text(file_path: str) -> str:
    """Extract all text from a .docx in document order, including tables."""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(file_path)
    lines: list[str] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # Plain paragraph — grab all run text
            text = "".join(r.text for r in block.iterchildren() if r.tag.endswith("}r"))
            # Also catch <w:t> directly
            if not text:
                text = "".join(t.text or "" for t in block.iter() if t.tag.endswith("}t"))
            text = text.strip()
            if text:
                lines.append(text)

        elif tag == "tbl":
            # Walk every cell in every row
            for row in block.iterchildren():  # <w:tr>
                if not row.tag.endswith("}tr"):
                    continue
                for cell in row.iterchildren():  # <w:tc>
                    if not cell.tag.endswith("}tc"):
                        continue
                    cell_text = "".join(
                        t.text or ""
                        for t in cell.iter()
                        if t.tag.endswith("}t")
                    ).strip()
                    if cell_text:
                        lines.append(cell_text)

    return "\n".join(lines)


class DocumentIngester:
    def __init__(self, index: SOPIndex, config: ChatbotConfig) -> None:
        self._index = index
        self._config = config

    def ingest(self, file_path: str) -> IngestResult:
        # Check extension
        _, ext = os.path.splitext(file_path)
        if ext.lower() not in SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            raise IngestError(
                f"Unsupported file format '{ext}'. Supported formats: {supported}"
            )

        # Read file
        try:
            if ext.lower() == ".docx":
                text = _extract_docx_text(file_path)
                # Extract, save, and describe images
                images_b64 = _extract_images_from_docx(file_path)
                image_paths = _save_images_from_docx(file_path)
                if images_b64:
                    image_descriptions = []
                    for i, img_b64 in enumerate(images_b64):
                        desc = _describe_image_with_vision(img_b64, i)
                        # Tag the description with the image path for display
                        img_path = image_paths[i] if i < len(image_paths) else ""
                        if img_path:
                            desc = f"[IMAGE_PATH:{img_path}]\n{desc}"
                        if desc:
                            image_descriptions.append(desc)
                    if image_descriptions:
                        text = text + "\n\n" + "\n\n".join(image_descriptions)
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
        except (FileNotFoundError, PermissionError, OSError) as e:
            raise IngestError(f"Cannot read file '{file_path}': {e}") from e

        # Split into chunks
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self._config.chunk_size,
            chunk_overlap=self._config.chunk_overlap,
        )
        raw_chunks = splitter.split_text(text)

        if not raw_chunks:
            raise IngestError(
                f"Document '{file_path}' produced no content after splitting."
            )

        file_name = os.path.basename(file_path)
        chunks = [
            Chunk(
                text=chunk_text,
                source=file_name,
                section_id=f"chunk_{i}",
                chunk_index=i,
            )
            for i, chunk_text in enumerate(raw_chunks)
        ]

        self._index.upsert_chunks(chunks)
        return IngestResult(file_name=file_name, chunk_count=len(chunks))

    def list_documents(self) -> list[str]:
        return self._index.list_sources()

    def remove_document(self, file_name: str) -> None:
        if file_name not in self.list_documents():
            raise DocumentNotFoundError(
                f"Document '{file_name}' not found in the index."
            )
        self._index.delete_by_source(file_name)

    def clear(self) -> None:
        self._index.clear()

    def get_chunks(self, file_name: str) -> list[Chunk]:
        result = self._index._collection.get(
            where={"source": file_name},
            include=["documents", "metadatas"],
        )
        chunks = [
            Chunk(
                text=doc,
                source=meta["source"],
                section_id=meta["section_id"],
                chunk_index=int(meta["chunk_index"]),
            )
            for doc, meta in zip(result["documents"], result["metadatas"])
        ]
        chunks.sort(key=lambda c: c.chunk_index)
        return chunks

    def pretty_print(self, file_name: str) -> str:
        chunks = self.get_chunks(file_name)
        return "\n\n".join(c.text for c in chunks)
